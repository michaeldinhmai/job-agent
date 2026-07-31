"""Resume reading and job-description gap analysis.

What this does: tells you which terms a job description leans on, and which of
them your resume already demonstrates. That's a *reporting* tool — it shows you
where your real experience isn't being surfaced in the language the posting
uses.

What it deliberately does not do: invent experience, or stuff keywords you
can't back up. Both are easy to detect and end the conversation when a recruiter
spots them. Rewrite bullets to describe things you actually did, using the
posting's vocabulary where it genuinely applies.
"""

from __future__ import annotations

import re
from collections import Counter
from pathlib import Path

STOPWORDS = {
    "the", "and", "for", "with", "you", "your", "our", "will", "are", "that",
    "this", "have", "from", "their", "them", "they", "who", "has", "was",
    "not", "but", "all", "can", "may", "any", "out", "how", "its", "his",
    "her", "she", "him", "each", "more", "most", "other", "into", "than",
    "then", "these", "those", "such", "only", "own", "same", "too", "very",
    "just", "also", "about", "over", "under", "across", "within", "while",
    "work", "working", "team", "teams", "role", "job", "position", "company",
    "years", "year", "experience", "ability", "strong", "excellent", "good",
    "new", "help", "make", "take", "well", "use", "using", "used", "including",
    "etc", "per", "via", "must", "should", "would", "could", "able", "like",
    "one", "two", "three", "five", "plus", "least", "based", "high", "level",
    "we", "us", "is", "in", "to", "of", "on", "as", "at", "by", "or", "an",
    "a", "be", "it", "if", "so", "do", "does", "what", "when", "where",
    "which", "you'll", "we're", "you're", "join", "apply", "please", "email",
    "nvidia", "llc", "inc", "equal", "opportunity", "employer", "benefits",
    "salary", "range", "compensation", "eligible", "applicants", "candidate",
    "candidates", "requirements", "qualifications", "responsibilities",
    "preferred", "required", "bachelor", "master", "degree", "equivalent",
    # salary/EEO boilerplate that every posting carries
    "usd", "base", "bonus", "equity", "status", "veteran", "disability",
    "race", "religion", "gender", "orientation", "through", "background",
    "collaborate", "collaborating", "passion", "passionate", "diverse",
    "world", "career", "mission", "impact", "growth", "environment",
}

# Terms worth flagging even if they appear only once — the vocabulary of this
# job family, where a single mention still signals a real requirement.
VOCAB = {
    "pre-sales", "presales", "post-sales", "poc", "proof of concept", "rfp",
    "rfi", "demo", "demonstration", "discovery", "technical win", "quota",
    "pipeline", "stakeholder", "enterprise", "saas", "customer-facing",
    "solution design", "architecture", "onboarding", "implementation",
    "migration", "integration", "api", "rest", "graphql", "sdk", "webhook",
    "python", "sql", "java", "javascript", "typescript", "go", "bash",
    "kubernetes", "docker", "terraform", "linux", "networking",
    "aws", "azure", "gcp", "cloud", "hybrid cloud", "on-prem",
    "snowflake", "databricks", "kafka", "spark", "airflow", "postgres",
    "observability", "monitoring", "security", "compliance", "sso", "saml",
    "llm", "genai", "generative ai", "rag", "inference", "fine-tuning",
    "machine learning", "deep learning", "pytorch", "tensorflow", "cuda",
    "gpu", "nim", "nemo", "triton", "vllm", "agentic", "multi-agent",
    "mlops", "data pipeline", "etl", "salesforce", "crm",
}

TOKEN_RE = re.compile(r"[a-z0-9][a-z0-9+#./-]*")


def read_docx(path: str | Path) -> str:
    """Extract all text from a .docx, including tables."""
    from docx import Document  # imported lazily so the rest of the CLI works without it

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(p for p in parts if p and p.strip())


def _tokens(text: str) -> list[str]:
    return TOKEN_RE.findall(text.lower())


def _phrases(text: str) -> Counter:
    """Count 1-, 2-, and 3-grams, dropping anything stopword-anchored."""
    words = _tokens(text)
    counts: Counter = Counter()
    for n in (1, 2, 3):
        for i in range(len(words) - n + 1):
            gram = words[i : i + n]
            if gram[0] in STOPWORDS or gram[-1] in STOPWORDS:
                continue
            if any(len(w) < 2 for w in gram):
                continue
            # pure numbers are salary figures, not skills
            if all(w.replace(",", "").replace(".", "").isdigit() for w in gram):
                continue
            counts[" ".join(gram)] += 1
    return counts


def has_term(term: str, text: str) -> bool:
    pattern = r"\b" + re.escape(term).replace(r"\ ", r"[\s-]") + r"\b"
    return re.search(pattern, text, re.I) is not None


def analyze(resume_text: str, jd_text: str, top: int = 30) -> dict:
    """Compare a resume against one job description.

    Returns terms the JD emphasises, split by whether the resume already
    evidences them.
    """
    counts = _phrases(jd_text)

    candidates: dict[str, int] = {}
    for term, n in counts.items():
        if n >= 2 or term in VOCAB:
            # Prefer the longer phrase when one contains another.
            candidates[term] = n * (2 if term in VOCAB else 1)

    ranked = sorted(candidates.items(), key=lambda kv: (-kv[1], kv[0]))
    matched, missing = [], []
    for term, weight in ranked:
        (matched if has_term(term, resume_text) else missing).append((term, weight))

    return {
        "matched": matched[:top],
        "missing": missing[:top],
        "coverage": len(matched) / max(len(ranked), 1),
        "jd_terms": len(ranked),
    }
